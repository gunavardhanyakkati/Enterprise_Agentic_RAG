import logging
from pathlib import Path

from src.exceptions import PDFParsingException
from src.schemas.pdf_parser.models import ParserType

from .models import ParsedDocument

logger = logging.getLogger(__name__)

SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt"}


class DocumentParserService:
    """Parse PDF, DOCX, and TXT uploads into plain text."""

    def validate_extension(self, filename: str) -> str:
        ext = Path(filename).suffix.lower()
        if ext not in SUPPORTED_EXTENSIONS:
            raise ValueError(
                f"Unsupported file type '{ext}'. Supported: {', '.join(sorted(SUPPORTED_EXTENSIONS))}"
            )
        return ext

    def parse(self, file_path: Path, filename: str | None = None) -> ParsedDocument:
        name = filename or file_path.name
        ext = self.validate_extension(name)
        file_size = file_path.stat().st_size

        if ext == ".pdf":
            return self._parse_pdf(file_path, name, file_size)
        if ext == ".docx":
            return self._parse_docx(file_path, name, file_size)
        return self._parse_txt(file_path, name, file_size)

    def _parse_pdf(self, file_path: Path, filename: str, file_size: int) -> ParsedDocument:
        try:
            return self._parse_pdf_pymupdf(file_path, filename, file_size)
        except Exception as pymupdf_error:
            logger.warning(f"PyMuPDF failed for {filename}: {pymupdf_error}. Trying PyPDF fallback.")
            try:
                return self._parse_pdf_pypdf(file_path, filename, file_size)
            except Exception as pypdf_error:
                raise PDFParsingException(
                    f"Failed to parse PDF '{filename}' with PyMuPDF and PyPDF: {pypdf_error}"
                ) from pypdf_error

    def _parse_pdf_pymupdf(self, file_path: Path, filename: str, file_size: int) -> ParsedDocument:
        import fitz

        page_texts: list[str] = []
        with fitz.open(file_path) as doc:
            for page in doc:
                page_texts.append(page.get_text("text"))

        content = "\n\n".join(text.strip() for text in page_texts if text.strip())
        return ParsedDocument(
            filename=filename,
            pages=len(page_texts) or 1,
            content=content,
            file_size_bytes=file_size,
            parser_used=ParserType.PYMUPDF,
            metadata={"parser": "pymupdf"},
            page_texts=page_texts,
        )

    def _parse_pdf_pypdf(self, file_path: Path, filename: str, file_size: int) -> ParsedDocument:
        from pypdf import PdfReader

        reader = PdfReader(str(file_path))
        page_texts = [(page.extract_text() or "") for page in reader.pages]
        content = "\n\n".join(text.strip() for text in page_texts if text.strip())

        return ParsedDocument(
            filename=filename,
            pages=len(page_texts) or 1,
            content=content,
            file_size_bytes=file_size,
            parser_used=ParserType.PYPDF,
            metadata={"parser": "pypdf"},
            page_texts=page_texts,
        )

    def _parse_docx(self, file_path: Path, filename: str, file_size: int) -> ParsedDocument:
        from docx import Document as DocxDocument

        doc = DocxDocument(str(file_path))
        paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        content = "\n\n".join(paragraphs)

        return ParsedDocument(
            filename=filename,
            pages=1,
            content=content,
            file_size_bytes=file_size,
            parser_used=ParserType.OFFICE,
            metadata={"parser": "python-docx", "paragraph_count": len(paragraphs)},
        )

    def _parse_txt(self, file_path: Path, filename: str, file_size: int) -> ParsedDocument:
        content = file_path.read_text(encoding="utf-8", errors="replace")
        return ParsedDocument(
            filename=filename,
            pages=1,
            content=content,
            file_size_bytes=file_size,
            parser_used=ParserType.TEXT,
            metadata={"parser": "text"},
        )
